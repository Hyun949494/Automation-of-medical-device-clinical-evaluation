# document_utils.py
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_word_document(content, pdf_length):
    """워드 문서 생성"""
    doc = Document()
    
    # 제목 추가
    title = doc.add_heading('Literature Analysis Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 분석 정보 추가
    doc.add_paragraph(f"Analysis Date: {pd.Timestamp.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Analysis Tool: Gemini AI-based Literature Analysis")
    doc.add_paragraph(f"Original Paper Text Length: {pdf_length} characters")
    
    # 구분선
    doc.add_paragraph("─" * 50)
    
    # 분석 내용 추가
    paragraphs = content.split('\n')
    for para in paragraphs:
        para = para.strip()
        if not para:
            continue
            
        if para.startswith('##'):
            level = min(para.count('#') - 1, 5)
            text = para.lstrip('#').strip()
            doc.add_heading(text, level)
        elif para.startswith('**') and para.endswith('**'):
            p = doc.add_paragraph()
            run = p.add_run(para.strip('*'))
            run.bold = True
        elif para.startswith('- '):
            doc.add_paragraph(para[2:], style='List Bullet')
        else:
            doc.add_paragraph(para)
    
    return doc